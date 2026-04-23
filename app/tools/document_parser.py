"""文档解析与切块工具。"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from app.schemas.qa import KnowledgeChunkIn


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    """解析后文档数据。"""

    title: str
    text: str


@dataclass(slots=True)
class DocumentParserTool:
    """支持 txt/md/pdf/docx 的解析与切块。"""

    chunk_size: int = 900
    overlap: int = 120

    def __post_init__(self) -> None:
        if self.chunk_size <= 0:
            raise ValueError("chunk_size must be greater than 0.")
        if self.overlap < 0:
            raise ValueError("overlap must be greater than or equal to 0.")
        if self.overlap >= self.chunk_size:
            raise ValueError("overlap must be less than chunk_size.")

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        if not content:
            raise ValueError("Uploaded file is empty.")

        suffix = Path(filename).suffix.lower()
        title = Path(filename).stem or "uploaded_document"

        if suffix in {".txt", ".md"}:
            text = content.decode("utf-8", errors="ignore")
            return ParsedDocument(title=title, text=self._normalize_text(text))

        if suffix == ".pdf":
            return ParsedDocument(title=title, text=self._parse_pdf(content))

        if suffix == ".docx":
            return ParsedDocument(title=title, text=self._parse_docx(content))

        raise ValueError("Unsupported file type. Allowed: .txt, .md, .pdf, .docx")

    def build_chunks(
        self,
        document_id: str,
        source_uri: str | None,
        parsed: ParsedDocument,
        ) -> list[KnowledgeChunkIn]:
        
        text = parsed.text.strip()
        if not text:
            raise ValueError("Parsed file is empty.")

        chunks: list[KnowledgeChunkIn] = []
        chunk_idx = 1

        for piece in self._split_text(text):
            chunks.append(
                KnowledgeChunkIn(
                    document_id=document_id,
                    chunk_id=f"chunk_{chunk_idx:04d}",
                    title=parsed.title,
                    content=piece,
                    source_uri=source_uri,
                )
            )
            chunk_idx += 1

        return chunks

    def _split_text(self, text: str) -> list[str]:
        if len(text) <= self.chunk_size:
            return [text]

        chunks: list[str] = []
        buffer = ""

        for segment in (segment.strip() for segment in text.split("\n")):
            if not segment:
                continue

            if len(segment) > self.chunk_size:
                if buffer:
                    chunks.extend(self._split_overlong_piece(buffer))
                    buffer = ""
                chunks.extend(self._split_overlong_piece(segment))
                continue

            candidate = segment if not buffer else f"{buffer}\n{segment}"
            if len(candidate) <= self.chunk_size:
                buffer = candidate
            else:
                chunks.extend(self._split_overlong_piece(buffer))
                buffer = segment

        if buffer:
            chunks.extend(self._split_overlong_piece(buffer))

        return chunks

    def _split_overlong_piece(self, text: str) -> list[str]:
        pieces: list[str] = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            piece = text[start:end].strip()
            if piece:
                pieces.append(piece)
            if end >= text_len:
                break
            start = max(0, end - self.overlap)

        return pieces

    @staticmethod
    def _normalize_text(text: str) -> str:
        return "\n".join(line.strip() for line in text.splitlines() if line.strip())

    def _parse_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Missing dependency `pypdf`. Please install it first.") from exc

        reader = PdfReader(BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return self._normalize_text("\n".join(pages))

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Missing dependency `python-docx`. Please install it first.") from exc

        doc = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return self._normalize_text(text)
